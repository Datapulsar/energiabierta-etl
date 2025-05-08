import os
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
from drive_functions import subir_csv_a_drive


def graficar_datos(df, ruta_imagen):
    """Genera un gráfico de emisiones de CO₂ y lo guarda como imagen."""
    print("Columnas del DataFrame:", df.columns)

    if 'Fecha' not in df.columns or 'Emisiones CO2 toneladas métricas per capita' not in df.columns:
        print("Error: Las columnas necesarias no están en el DataFrame.")
        return

    print("Primeras filas del DataFrame:")
    print(df.head())

    plt.figure(figsize=(10, 6))
    plt.plot(df['Fecha'], df['Emisiones CO2 toneladas métricas per capita'])
    plt.title('Emisiones de CO₂ per cápita')
    plt.xlabel('Fecha')
    plt.ylabel('Emisiones CO₂ (toneladas métricas per cápita)')
    plt.grid(True)
    plt.savefig(ruta_imagen)
    plt.close()


def generar_informe(df, carpeta_drive_id, service):
    # Calcular resumen estadístico
    resumen_stats = df['Emisiones CO2 toneladas métricas per capita'].describe().to_string()

    # Mostrar en consola
    print("\nResumen estadístico de las emisiones de CO₂:")
    print(resumen_stats)

    # Mostrar gráfico en pantalla
    try:
        df.plot(x='Fecha', y='Emisiones CO2 toneladas métricas per capita',
                title='Emisiones de CO₂ per cápita', figsize=(10, 6), grid=True)
        plt.xlabel('Fecha')
        plt.ylabel('Emisiones CO₂ (toneladas métricas per cápita)')
        plt.tight_layout()
        plt.show()
    except Exception as e:
        print(f"Error al mostrar el gráfico: {e}")

    # Definir rutas
    ruta_pdf = "data/informe_emisiones.pdf"
    ruta_doc = "data/informe_emisiones.docx"
    ruta_imagen = "data/grafico_emisiones.png"

    # Crear imagen del gráfico
    try:
        graficar_datos(df, ruta_imagen)
        print(f"Gráfico generado en: {ruta_imagen}")
    except Exception as e:
        print(f"Error al generar el gráfico: {e}")
        return

    # Crear informes
    try:
        generar_informe_pdf(ruta_pdf, ruta_imagen, resumen_stats)
    except Exception as e:
        print(f"Error al generar el informe PDF: {e}")
        return

    try:
        generar_informe_doc(ruta_doc, ruta_imagen, resumen_stats)
    except Exception as e:
        print(f"Error al generar el informe DOC: {e}")
        return

    # Subir a Google Drive
    try:
        subir_csv_a_drive(service, ruta_pdf, carpeta_drive_id)
        subir_csv_a_drive(service, ruta_doc, carpeta_drive_id)
    except Exception as e:
        print(f"Error al subir los archivos a Google Drive: {e}")


def generar_informe_pdf(ruta_pdf, ruta_imagen, resumen_stats):
    """Genera un informe en formato PDF."""
    try:
        c = canvas.Canvas(ruta_pdf, pagesize=letter)
        c.drawString(100, 750, "Informe de Análisis de Emisiones de CO₂")
        c.drawString(100, 730, "Análisis estadístico y gráfico generado a partir de los datos.")

        # Añadir resumen estadístico
        c.drawString(100, 700, "Resumen estadístico:")
        y = 680
        for linea in resumen_stats.split('\n'):
            c.drawString(100, y, linea)
            y -= 15
            if y < 100:
                c.showPage()
                y = 750

        # Incluir imagen si hay espacio suficiente
        if os.path.exists(ruta_imagen):
            if y < 400:
                c.showPage()
                y = 700
            c.drawImage(ruta_imagen, 100, y - 320, width=400, height=300)

        c.showPage()
        c.save()
        print(f"Informe PDF guardado en: {ruta_pdf}")
    except Exception as e:
        print(f"Error al generar el informe PDF: {e}")


def generar_informe_doc(ruta_doc, ruta_imagen, resumen_stats):
    """Genera un informe en formato DOCX."""
    try:
        doc = Document()
        doc.add_heading('Informe de Análisis de Emisiones de CO₂', 0)

        doc.add_paragraph(
            "Este informe incluye un análisis estadístico y gráfico sobre las emisiones de CO₂ per cápita a lo largo de los años."
        )

        doc.add_heading("Resumen estadístico:", level=1)
        doc.add_paragraph(resumen_stats)

        if os.path.exists(ruta_imagen):
            doc.add_picture(ruta_imagen, width=Inches(5))

        doc.save(ruta_doc)
        print(f"Informe DOC guardado en: {ruta_doc}")
    except Exception as e:
        print(f"Error al generar el informe DOC: {e}")
